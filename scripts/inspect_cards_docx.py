from docx import Document

from yiju_paths import find_cards_document

path = find_cards_document()
print("Using", path)
doc = Document(str(path))

print("--- PARAGRAPHS (non-empty, first 100) ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(i, t[:100])

print("--- TABLES", len(doc.tables), "---")
for ti, table in enumerate(doc.tables):
    print("TABLE", ti, "rows", len(table.rows), "cols", len(table.columns))
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip().replace("\n", " ")[:50] for c in row.cells]
        print(f"  {ri}: {cells}")
